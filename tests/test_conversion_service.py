"""
Testes unitários para ConversionService.

História de Usuário:
  Como pipeline de processamento,
  Quero converter HTMLs traduzidos para .docx e .pdf,
  Para entregar os documentos em formatos amplamente utilizados.
"""
from pathlib import Path

import pytest
from docx import Document

from backend.services.conversion_service import ConversionService

SAMPLE_HTML = """<!DOCTYPE html>
<html><head><title>Documento de Teste</title></head>
<body>
<h1>Título do Documento</h1>
<h2>Introdução</h2>
<p>Este é um parágrafo com conteúdo relevante para o teste.</p>
<h3>Métodos</h3>
<p>Os procedimentos foram aplicados conforme o padrão estabelecido.</p>
<ul>
  <li>Item 1 da lista</li>
  <li>Item 2 da lista</li>
</ul>
</body></html>"""


def test_html_to_docx_creates_file(tmp_path: Path) -> None:
    """Deve criar o arquivo .docx no caminho especificado."""
    output = tmp_path / "output.docx"

    ConversionService().html_to_docx(SAMPLE_HTML, output)

    assert output.exists()
    assert output.stat().st_size > 0


def test_html_to_docx_produces_valid_docx(tmp_path: Path) -> None:
    """O arquivo gerado deve ser um .docx válido e legível."""
    output = tmp_path / "output.docx"

    ConversionService().html_to_docx(SAMPLE_HTML, output)
    doc = Document(str(output))

    texts = [p.text for p in doc.paragraphs if p.text]
    assert len(texts) > 0
    assert any("Título" in t for t in texts)


def test_html_to_docx_includes_headings(tmp_path: Path) -> None:
    """Deve incluir os headings do HTML como parágrafos de heading no docx."""
    output = tmp_path / "output.docx"

    ConversionService().html_to_docx(SAMPLE_HTML, output)
    doc = Document(str(output))

    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(heading_texts) > 0


def test_html_to_pdf_creates_file(tmp_path: Path) -> None:
    """Deve criar o arquivo .pdf no caminho especificado."""
    output = tmp_path / "output.pdf"

    ConversionService().html_to_pdf(SAMPLE_HTML, output)

    assert output.exists()
    assert output.stat().st_size > 0


def test_html_to_pdf_starts_with_pdf_header(tmp_path: Path) -> None:
    """O arquivo PDF gerado deve começar com o header %PDF."""
    output = tmp_path / "output.pdf"

    ConversionService().html_to_pdf(SAMPLE_HTML, output)

    assert output.read_bytes()[:4] == b"%PDF"


def test_html_to_markdown_creates_file(tmp_path: Path) -> None:
    """Deve criar o arquivo .md no caminho especificado."""
    output = tmp_path / "output.md"

    ConversionService().html_to_markdown(SAMPLE_HTML, output)

    assert output.exists()
    assert output.stat().st_size > 0


def test_html_to_markdown_produces_valid_markdown(tmp_path: Path) -> None:
    """O arquivo gerado deve conter a marcação markdown correspondente aos elementos HTML."""
    output = tmp_path / "output.md"

    ConversionService().html_to_markdown(SAMPLE_HTML, output)
    content = output.read_text(encoding="utf-8")

    assert "# Título do Documento" in content
    assert "## Introdução" in content
    assert "### Métodos" in content
    assert "- Item 1 da lista" in content
    assert "- Item 2 da lista" in content
    assert "Este é um parágrafo" in content


def test_html_to_markdown_extracts_base64_images(tmp_path: Path) -> None:
    """Deve extrair imagens Base64 inline para arquivos físicos na pasta 'images'."""
    # Um HTML com uma tag img contendo imagem inline em Base64 (pequena imagem PNG válida de 1 pixel)
    base64_img = (
        "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk"
        "+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
    )
    html_with_img = f'<html><body><p>Texto antes</p><img alt="Figura Teste" src="{base64_img}" /><p>Texto depois</p></body></html>'

    output = tmp_path / "artigo.md"
    ConversionService().html_to_markdown(html_with_img, output)

    # 1. O arquivo markdown deve existir
    assert output.exists()
    content = output.read_text(encoding="utf-8")

    # 2. A referência da imagem no markdown deve ser relativa e apontar para ./images/artigo_img_1.png
    assert "![Figura Teste](./images/artigo_img_1.png)" in content

    # 3. A pasta 'images' deve ter sido criada ao lado do markdown
    images_dir = tmp_path / "images"
    assert images_dir.exists()
    assert images_dir.is_dir()

    # 4. A imagem física deve ter sido gravada no disco
    target_img = images_dir / "artigo_img_1.png"
    assert target_img.exists()
    assert target_img.stat().st_size > 0
