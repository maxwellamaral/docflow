"""Conversão de HTML para os formatos Word (.docx) e PDF."""
import io
import re
import base64
from pathlib import Path

import weasyprint
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches


class ConversionService:
    """Serviço de conversão de HTML para .docx e .pdf."""

    def html_to_docx(self, html_content: str, output_path: Path) -> Path:
        """Converte um HTML para formato Word (.docx).

        Extrai headings, parágrafos, imagens e itens de lista do HTML e os
        mapeia para os estilos correspondentes do python-docx.

        Args:
            html_content: Conteúdo HTML de entrada.
            output_path: Caminho de destino para o arquivo .docx.

        Returns:
            Caminho do arquivo .docx gerado.
        """
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "img"]):
            if element.name == "img":
                src = element.get("src", "")
                if src.startswith("data:image/"):
                    try:
                        header, encoded = src.split(",", 1)
                        img_bytes = base64.b64decode(encoded)
                        doc.add_picture(io.BytesIO(img_bytes), width=Inches(5))
                    except Exception:
                        pass
                continue

            text = element.get_text(strip=True)
            if not text:
                continue

            tag = element.name
            if tag == "h1":
                doc.add_heading(text, level=1)
            elif tag == "h2":
                doc.add_heading(text, level=2)
            elif tag in ("h3", "h4"):
                doc.add_heading(text, level=3)
            elif tag in ("h5", "h6"):
                doc.add_heading(text, level=4)
            elif tag == "li":
                try:
                    doc.add_paragraph(text, style="List Bullet")
                except KeyError:
                    doc.add_paragraph(f"• {text}")
            else:
                doc.add_paragraph(text)

        doc.save(output_path)
        return output_path

    def html_to_pdf(self, html_content: str, output_path: Path) -> Path:
        """Converte um HTML para PDF usando WeasyPrint.

        Args:
            html_content: Conteúdo HTML de entrada.
            output_path: Caminho de destino para o arquivo .pdf.

        Returns:
            Caminho do arquivo .pdf gerado.
        """
        weasyprint.HTML(string=html_content).write_pdf(str(output_path))
        return output_path

    def html_to_markdown(self, html_content: str, output_path: Path) -> Path:
        """Converte um HTML para formato Markdown (.md).

        Navega pelo DOM do HTML e gera o Markdown equivalente, extraindo
        imagens Base64 para uma pasta física 'images' ao lado do arquivo .md,
        referenciando-as com caminhos relativos.

        Args:
            html_content: Conteúdo HTML de entrada.
            output_path: Caminho de destino para o arquivo .md.

        Returns:
            Caminho do arquivo .md gerado.
        """
        soup = BeautifulSoup(html_content, "html.parser")

        # Cria a pasta 'images' dentro da pasta de destino do markdown
        markdown_dir = output_path.parent
        images_dir = markdown_dir / "images"
        images_dir.mkdir(parents=True, exist_ok=True)

        # Nome base para evitar colisão de imagens de múltiplos arquivos convertidos
        file_stem = output_path.stem
        img_counter = [0]  # Usamos uma lista mutável para atuar como contador no escopo aninhado

        def _parse_element(element, is_nested=False) -> str:
            if element.name is None:
                return element.string or ""

            name = element.name

            if name in ["html", "body", "div", "section", "article"]:
                parts = []
                for child in element.children:
                    parts.append(_parse_element(child, is_nested))
                return "".join(parts)

            elif name == "h1":
                return f"\n\n# {element.get_text(strip=True)}\n\n"
            elif name == "h2":
                return f"\n\n## {element.get_text(strip=True)}\n\n"
            elif name == "h3":
                return f"\n\n### {element.get_text(strip=True)}\n\n"
            elif name == "h4":
                return f"\n\n#### {element.get_text(strip=True)}\n\n"
            elif name in ["h5", "h6"]:
                return f"\n\n##### {element.get_text(strip=True)}\n\n"
            elif name == "p":
                inner_text = "".join(_parse_element(c, is_nested) for c in element.children)
                return f"\n\n{inner_text.strip()}\n\n"
            elif name == "ul":
                parts = []
                for child in element.children:
                    if child.name == "li":
                        parts.append(f"\n- {_parse_element(child, True).strip()}")
                    else:
                        parts.append(_parse_element(child, is_nested))
                return "".join(parts) + "\n"
            elif name == "ol":
                parts = []
                idx = 1
                for child in element.children:
                    if child.name == "li":
                        parts.append(f"\n{idx}. {_parse_element(child, True).strip()}")
                        idx += 1
                    else:
                        parts.append(_parse_element(child, is_nested))
                return "".join(parts) + "\n"
            elif name == "li":
                return "".join(_parse_element(c, is_nested) for c in element.children)
            elif name in ["strong", "b"]:
                return f" **{element.get_text()}** "
            elif name in ["em", "i"]:
                return f" *{element.get_text()}* "
            elif name == "a":
                href = element.get("href", "")
                text = element.get_text()
                return f" [{text}]({href}) "
            elif name == "img":
                alt = element.get("alt", "imagem")
                src = element.get("src", "")

                # Se for imagem inline Base64, decodifica e salva em arquivo físico
                if src.startswith("data:image/"):
                    try:
                        mime_header, encoded_data = src.split(",", 1)
                        ext = "png"
                        if "image/" in mime_header:
                            parts = mime_header.split(";")[0].split("/")
                            if len(parts) > 1:
                                ext = parts[1]

                        img_bytes = base64.b64decode(encoded_data)
                        img_counter[0] += 1

                        image_filename = f"{file_stem}_img_{img_counter[0]}.{ext}"
                        image_file_path = images_dir / image_filename
                        image_file_path.write_bytes(img_bytes)

                        return f"\n\n![{alt}](./images/{image_filename})\n\n"
                    except Exception:
                        pass

                return f"\n\n![{alt}]({src})\n\n"
            elif name == "br":
                return "\n"
            else:
                return "".join(_parse_element(c, is_nested) for c in element.children)

        body = soup.find("body") or soup
        md_text = _parse_element(body)
        md_text = re.sub(r"\n{3,}", "\n\n", md_text)

        output_path.write_text(md_text.strip() + "\n", encoding="utf-8")
        return output_path
